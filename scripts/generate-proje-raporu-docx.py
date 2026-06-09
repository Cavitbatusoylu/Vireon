# -*- coding: utf-8 -*-
"""Vireon 15 sayfalık proje raporunu Word (.docx) olarak üretir."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Vireon-Proje-Raporu.docx")


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "E8F4F8")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 80, 120)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def build_document():
    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- KAPAK ---
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("VIREON")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0, 140, 180)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Dijital Banka Çekirdek Sistemi Simülasyonu\nProje Raporu")
    r2.font.size = Pt(16)
    r2.bold = True

    doc.add_paragraph()
    cover_lines = [
        ("Proje Türü", "2026 Final Projesi — Eğitim Amaçlı"),
        ("Teknoloji", "ASP.NET Core 8 · EF Core · SQLite · PWA"),
        ("Ekip", "Cavit Batu Soylu · Enes Kaya · Kerem Arslan"),
        ("Repository", "github.com/Cavitbatusoylu/Vireon"),
        ("Rapor Tarihi", "Haziran 2026"),
    ]
    for label, val in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        p.add_run(val)

    doc.add_page_break()

    # --- ÖZET ---
    add_heading(doc, "Özet (Abstract)", 1)
    add_para(
        doc,
        "Vireon, gerçek bir banka çekirdek sisteminin temel bileşenlerini simüle eden, veritabanı "
        "merkezli bir dijital bankacılık uygulamasıdır. Proje; ACID uyumlu para transferi, para yatırma, "
        "değiştirilemez muhasebe defteri (immutable ledger), kural tabanlı fraud tespiti, günlük transfer "
        "limiti yönetimi ve yapay zeka destekli müşteri asistanı (Neon AI) modüllerini tek bir platformda "
        "birleştirir."
    )
    add_para(
        doc,
        "Sistem, beş katmanlı N-Tier mimari üzerine inşa edilmiştir: Entity, Data Access, DTO, Business "
        "ve Presentation. Veri kalıcılığı SQLite ile sağlanır; ekip çalışması için paylaşımlı veritabanı "
        "dosyası Git tabanlı otomatik senkronizasyon ile yönetilir. Sunum katmanı hem RESTful Web API "
        "hem de PWA destekli modern web arayüzünden oluşur."
    )
    add_para(
        doc,
        "Anahtar Kelimeler: Dijital bankacılık, ACID, immutable ledger, fraud detection, SQLite, "
        "ASP.NET Core, PWA, yapay zeka asistanı",
        italic=True,
    )

    # --- 1. GİRİŞ ---
    add_heading(doc, "1. Giriş", 1)
    add_heading(doc, "1.1 Problem Tanımı", 2)
    add_para(
        doc,
        "Geleneksel bankacılık sistemleri; hesap yönetimi, para transferi, muhasebe kaydı, risk analizi "
        "ve müşteri iletişimi gibi birbirine bağımlı modüllerden oluşur. Üniversite final projelerinde "
        "genellikle ya yalnızca arayüz ya da yalnızca API geliştirilir; oysa gerçek dünyada finansal "
        "sistemler uçtan uca entegre çalışır. Vireon bu boşluğu doldurmak amacıyla tasarlanmıştır."
    )
    add_heading(doc, "1.2 Proje Hedefleri", 2)
    add_bullets(
        doc,
        [
            "Finansal tutarlılık: Transfer ve yatırma işlemlerinde bakiye, işlem kaydı ve defter girişlerinin atomik güncellenmesi",
            "İzlenebilirlik: Her bakiye değişiminin LedgerEntries tablosunda PreviousBalance / NewBalance ile kayıt altına alınması",
            "Risk yönetimi: Kural tabanlı fraud motoru ile şüpheli işlemlerin engellenmesi ve loglanması",
            "Kullanılabilirlik: TR/EN dil desteği, responsive PWA arayüzü, admin paneli",
            "Ekip çalışması: Paylaşımlı SQLite + Git senkronu",
            "Yapay zeka entegrasyonu: Groq API ile Neon AI asistanı (offline fallback destekli)",
        ],
    )
    add_heading(doc, "1.3 Proje Kapsamı", 2)
    add_para(doc, "Kapsam dahilinde: kullanıcı yönetimi, transfer/yatırma, limit, işlem geçmişi, fraud logları, Neon AI, Docker deploy.", bold=False)
    add_para(doc, "Kapsam dışında: gerçek ödeme ağ geçitleri, enterprise JWT/OAuth, otomatik test suite, production KVKK/PCI denetimi.", bold=False)

    # --- 2. MİMARİ ---
    add_heading(doc, "2. Sistem Mimarisi", 1)
    add_heading(doc, "2.1 Katmanlı Mimari (N-Tier)", 2)
    add_table(
        doc,
        ["Katman", "Proje", "Sorumluluk"],
        [
            ("Entity", "Vireon.EntityLayer", "Domain modelleri (User, Account, Transaction…)"),
            ("Data Access", "Vireon.DataAccessLayer", "EF Core DbContext, migration'lar"),
            ("DTO", "Vireon.DtoLayer", "API request/response modelleri"),
            ("Business", "Vireon.BusinessLayer", "TransactionManager, FraudModelService, NeonAIService"),
            ("Presentation", "Vireon.PresentationLayer", "Web API, middleware, wwwroot UI"),
        ],
        col_widths=[3, 5, 7],
    )
    add_para(doc, "Bağımlılık yönü: Presentation → Business + Data Access + DTO → Entity. Alt katmanlar üst katmanları tanımaz.")

    add_heading(doc, "2.2 Çalışma Zamanı Bileşenleri", 2)
    add_bullets(
        doc,
        [
            "Serilog yapılandırması (console + dosya log)",
            "SQLite bağlantısı (Database/vireon_local.db — sabit yol çözümü)",
            "DI kayıtları: DbContext, AutoMapper, FluentValidation, ITransactionService, FraudModelService, NeonAIService",
            "GitHub'dan DB pull (opsiyonel) ve otomatik migration",
            "Demo hesap seed (EnsureDemoAccounts)",
            "Middleware pipeline: Swagger, ErrorHandling, güvenlik başlıkları, CORS, static files",
        ],
    )

    add_heading(doc, "2.3 Transfer İstek Akışı", 2)
    add_para(
        doc,
        "Kullanıcı UI → POST /api/transfers/send-by-account → TransfersController → FraudModelService.Predict() "
        "→ (risk > 0.70 ise engelle) → TransactionManager.ProcessTransaction() → DB transaction (bakiye, ledger, fraud log) "
        "→ UI fetchDashboardData() ile güncel bakiye."
    )

    # --- 3. TEKNOLOJİ ---
    add_heading(doc, "3. Teknoloji Yığını", 1)
    add_table(
        doc,
        ["Kategori", "Teknoloji", "Kullanım"],
        [
            ("Backend", ".NET 8 / ASP.NET Core", "Web API + static hosting"),
            ("ORM", "Entity Framework Core", "Code-First migration"),
            ("Veritabanı", "SQLite", "Geliştirme ve demo"),
            ("Güvenlik", "BCrypt.Net", "Şifre hash"),
            ("Doğrulama", "FluentValidation", "DTO kuralları"),
            ("Mapping", "AutoMapper", "DTO → Entity"),
            ("Log", "Serilog", "Yapılandırılmış log"),
            ("Frontend", "HTML/CSS/JS", "SPA benzeri PWA"),
            ("Grafik", "Chart.js", "Dashboard"),
            ("AI", "Groq API (llama-3.1-8b-instant)", "Neon AI"),
            ("Deploy", "Docker, Electron", "Container ve masaüstü"),
        ],
        col_widths=[3, 5, 7],
    )

    # --- 4. VERİTABANI ---
    add_heading(doc, "4. Veritabanı Tasarımı", 1)
    add_heading(doc, "4.1 Tablolar", 2)
    add_table(
        doc,
        ["Tablo", "Açıklama", "Kritik Alanlar"],
        [
            ("Users", "Kullanıcılar", "Email (unique), Password (BCrypt), AccountNumber, Role"),
            ("Accounts", "Hesaplar", "Balance, Currency, RowVersion (concurrency)"),
            ("Transactions", "Transfer/yatırma", "SenderAccountId, ReceiverAccountId, Status"),
            ("LedgerEntries", "Immutable defter", "PreviousBalance, NewBalance, Amount"),
            ("DailyLimits", "Günlük limit", "MaxDailyLimit, UsedLimit, LastResetDate"),
            ("FraudLogs", "Risk kayıtları", "RiskType, Description, LogDate"),
        ],
        col_widths=[3.5, 4, 7],
    )

    add_heading(doc, "4.2 İlişkiler ve Kısıtlar", 2)
    add_bullets(
        doc,
        [
            "Transaction → Account: iki ayrı FK (Sender, Receiver), DeleteBehavior.Restrict",
            "DailyLimit ↔ User: birebir ilişki",
            "Tüm decimal alanlar decimal(18,2)",
            "Performans indexleri: CreatedAt, SenderAccountId, AccountId",
            "Account.RowVersion: eşzamanlı transfer koruması (DbUpdateConcurrencyException)",
        ],
    )

    add_heading(doc, "4.3 Migration Geçmişi", 2)
    add_table(
        doc,
        ["Migration", "İçerik"],
        [
            ("InitialCreate", "Temel 6 tablo"),
            ("AddRowVersion", "Account concurrency token"),
            ("AddUserRole", "Admin/User rolü"),
            ("AddPlainPassword", "Demo okunabilir şifre sütunu"),
        ],
        col_widths=[5, 10],
    )

    add_heading(doc, "4.4 Demo Seed Verisi", 2)
    add_table(
        doc,
        ["Kullanıcı", "E-posta", "Hesap", "Bakiye", "Rol"],
        [
            ("Cavit Batu Soylu", "cavit@vireon.com", "VR-99999", "1.000.000 ₺", "Admin"),
            ("Enes Kaya", "enes@vireon.com", "VR-88888", "50.000 ₺", "User"),
            ("Kerem Arslan", "kerem@vireon.com", "VR-77777", "50.000 ₺", "User"),
        ],
        col_widths=[4, 4.5, 2.5, 3, 2.5],
    )

    doc.add_page_break()

    # --- 5. API ---
    add_heading(doc, "5. Backend API Katmanı", 1)
    add_table(
        doc,
        ["Controller", "Endpoint'ler", "Sorumlu"],
        [
            ("UsersController", "register, login, forgot-password, delete-account, admin-*", "Cavit"),
            ("TransfersController", "send, send-by-account, deposit", "Enes"),
            ("AccountsController", "CRUD hesap", "Enes"),
            ("TransactionsController", "İşlem listeleme", "Enes"),
            ("LedgerEntriesController", "Defter kayıtları", "Enes"),
            ("FraudLogsController", "Fraud log listeleme", "Enes"),
            ("DailyLimitsController", "Limit CRUD", "Enes"),
            ("AIController", "POST /api/ai/chat", "Cavit"),
        ],
        col_widths=[4, 7, 3],
    )

    add_heading(doc, "5.1 UsersController", 2)
    add_para(doc, "Kayıt atomik transaction ile User + Account + DailyLimit + Ledger oluşturur. Giriş BCrypt doğrulama yapar. Şifre sıfırlama e-posta + hesap no eşleşmesi gerektirir. Hesap silme cascade ile tüm ilişkili veriyi temizler; admin demo hesabı korunur.")

    add_heading(doc, "5.2 ErrorHandlingMiddleware", 2)
    add_para(doc, "InvalidOperationException → 400 BUSINESS_RULE_VIOLATION; KeyNotFoundException → 404; diğer → 500 INTERNAL_ERROR. Güvenlik başlıkları: X-Content-Type-Options, X-Frame-Options, X-XSS-Protection.")

    # --- 6. İŞ KURALLARI ---
    add_heading(doc, "6. İş Kuralları ve Finansal Mantık", 1)
    add_heading(doc, "6.1 TransactionManager", 2)
    add_bullets(
        doc,
        [
            "Gönderici/alıcı hesap kontrolü; kendine transfer yasak; tutar > 0",
            "Döviz çevrimi (TRY/USD/EUR/GBP sabit kurlar)",
            "Yetersiz bakiye kontrolü",
            "Günlük limit — aşımda FraudLog (LIMIT_EXCEEDED) + exception",
            "Bakiye güncelleme + Transaction (Completed) + 2 LedgerEntry",
            "Fraud log: HIGH_AMOUNT (>10k), NIGHT (>5k, 00–05), FREQUENT (3+/dk)",
            "DbUpdateConcurrencyException → rollback + kullanıcıya tekrar dene mesajı",
        ],
    )

    add_heading(doc, "6.2 ACID Özellikleri", 2)
    add_table(
        doc,
        ["Özellik", "Uygulama"],
        [
            ("Atomicity", "BeginTransaction / Commit / Rollback"),
            ("Consistency", "Bakiye + ledger aynı transaction'da"),
            ("Isolation", "SQLite transaction izolasyonu"),
            ("Durability", "SQLite dosyaya kalıcı yazım"),
        ],
        col_widths=[4, 11],
    )

    # --- 7. FRAUD ---
    add_heading(doc, "7. Fraud Detection Sistemi", 1)
    add_heading(doc, "7.1 FraudModelService (Engelleme Katmanı)", 2)
    add_para(
        doc,
        "Üç özellik: amount (tutar/25000), hour (gece 0–5 veya 23+), frequency (işlem sayısı/8). "
        "Skor = 0.6×amount + 0.2×night + 0.2×frequency. Eşik: 0.70 (%70). Üzerinde transfer engellenir."
    )
    add_heading(doc, "7.2 TransactionManager (Loglama Katmanı)", 2)
    add_table(
        doc,
        ["RiskType", "Tetikleyici"],
        [
            ("LIMIT_EXCEEDED", "Günlük limit aşım denemesi"),
            ("HIGH_AMOUNT", "Transfer > 10.000 ₺"),
            ("SUSPICIOUS_NIGHT_TRANSFER", "00:00–05:59, tutar > 5.000"),
            ("FREQUENT_TRANSACTIONS", "1 dakikada 3+ işlem"),
        ],
        col_widths=[5.5, 9],
    )

    # --- 8. NEON AI ---
    add_heading(doc, "8. Neon AI — Yapay Zeka Asistanı", 1)
    add_para(doc, "NeonAIService + AIController (POST /api/ai/chat). Groq API llama-3.1-8b-instant modeli. API anahtarı yoksa veya hata olursa offline TR/EN sözlük fallback. Frontend floating chat + dashboard sohbet; localStorage ile geçmiş (son 10 mesaj API'ye gönderilir).")

    # --- 9. FRONTEND ---
    add_heading(doc, "9. Frontend ve Kullanıcı Deneyimi", 1)
    add_bullets(
        doc,
        [
            "index.html (~1700 satır): landing, modals, dashboard, admin, Neon AI",
            "vireon.js (~2400 satır): auth, fetchDashboardData, transfer, admin, i18n",
            "vireon.css (~5400 satır): neon tema, responsive, floating chat",
            "TR/EN dil desteği (data-tr / data-en)",
            "Dark/Light tema, PWA manifest + service worker",
            "Session: localStorage; rate limit 5/dk; idle timeout 15 dk",
            "Dashboard gerçek DB bakiyesi; admin panel sunum bakiyeleri (100k/50k)",
        ],
    )

    doc.add_page_break()

    # --- 10. GÜVENLİK ---
    add_heading(doc, "10. Güvenlik", 1)
    add_bullets(
        doc,
        [
            "BCrypt şifre hash; PlainPassword yalnızca demo amaçlı",
            "FluentValidation input doğrulama",
            "XSS: escapeHtml() + güvenlik response header'ları",
            "RowVersion concurrency; DB transaction ile atomik bakiye",
            "Bilinen sınırlama: API endpoint'lerinde JWT/auth middleware yok (demo kapsamı)",
        ],
    )

    # --- 11. DB SYNC ---
    add_heading(doc, "11. Paylaşımlı Veritabanı ve Ekip Senkronizasyonu", 1)
    add_para(doc, "SharedDatabaseGitSync: PullOnStartup ile GitHub'dan DB çekme; FileSystemWatcher ile değişiklik sonrası otomatik commit/push (~8 sn). ResolveSharedDbPath() cwd'den bağımsız Database/vireon_local.db kullanır. Manuel: npm run sync-db.")

    # --- 12. EKİP ---
    add_heading(doc, "12. Ekip ve Görev Dağılımı", 1)
    add_table(
        doc,
        ["Kişi", "Ana Alan", "Sorumluluklar"],
        [
            ("Cavit Batu Soylu", "Full Stack / UI", "PWA, Neon AI, UsersController, Git DB sync, deploy, README"),
            ("Enes Kaya", "API / İş Kuralları", "Finans API'leri, TransactionManager, FraudModelService, Validation"),
            ("Kerem Arslan", "Veritabanı", "Entity, VireonContext, migration, demo seed"),
        ],
        col_widths=[3.5, 3.5, 8.5],
    )
    add_para(doc, "Repository 117 commit içerir. Önemli kilometre taşları: N-Tier mimari, SQLite geçişi, ACID transfer, fraud motoru, Neon AI, paylaşımlı DB, profil yönetimi.")

    # --- 13. KURULUM ---
    add_heading(doc, "13. Kurulum, Çalıştırma ve Dağıtım", 1)
    add_para(doc, "git clone → cd Vireon.PresentationLayer → dotnet run → http://localhost:5202. Docker: docker build -t vireon . && docker run -p 5202:5202 vireon. Loglar: logs/vireon-YYYYMMDD.log (Serilog).")

    # --- 14. TEST ---
    add_heading(doc, "14. Test ve Doğrulama", 1)
    add_table(
        doc,
        ["#", "Senaryo", "Beklenen Sonuç"],
        [
            ("1", "Yeni kullanıcı kaydı", "VR-XXXXX hesap, 0 bakiye, 50k limit"),
            ("2", "Demo hesapla giriş", "Dashboard, doğru bakiye"),
            ("3", "500 ₺ transfer", "Bakiye düşer, ledger + transaction"),
            ("4", "Yetersiz bakiye", "Hata, bakiye değişmez"),
            ("5", "30.000 ₺ transfer", "AI blocked"),
            ("6", "Neon 'yardım'", "Offline rehber yanıtı"),
            ("7", "Admin panel", "Kullanıcı + fraud listesi"),
            ("8", "Şifre sıfırlama", "Yeni şifreyle giriş"),
            ("9", "git pull + dotnet run", "Paylaşımlı işlem geçmişi"),
            ("10", "PWA install", "Manifest + SW kayıtlı"),
        ],
        col_widths=[1, 5.5, 8],
    )

    # --- 15. SONUÇ ---
    add_heading(doc, "15. Sonuç ve Gelecek Çalışmalar", 1)
    add_heading(doc, "15.1 Elde Edilen Sonuçlar", 2)
    add_para(
        doc,
        "Vireon, dijital bankacılık çekirdek sisteminin temel bileşenlerini eğitim ortamında başarıyla "
        "simüle etmektedir. ACID transferler, immutable ledger, fraud tespiti, limit yönetimi, modern "
        "PWA arayüzü ve yapay zeka asistanı tek entegre platformda sunulmuştur."
    )
    add_heading(doc, "15.2 Gelecek İyileştirmeler", 2)
    add_bullets(
        doc,
        [
            "JWT tabanlı authentication ve role-based authorization",
            "Unit / integration testleri",
            "AI engellenen işlemlerin Failed status ile history'de görünmesi",
            "SignalR ile gerçek zamanlı bildirimler",
            "Production DB (PostgreSQL / SQL Server)",
            "KVKK uyumu — PlainPassword kaldırma",
            "CI/CD pipeline (GitHub Actions)",
        ],
    )

    # --- EKLER ---
    add_heading(doc, "Ek A — API Endpoint Özeti", 1)
    add_table(
        doc,
        ["Method", "Endpoint", "Açıklama"],
        [
            ("POST", "/api/users/register", "Kayıt"),
            ("POST", "/api/users/login", "Giriş"),
            ("POST", "/api/users/forgot-password", "Şifre sıfırlama"),
            ("POST", "/api/users/{id}/delete-account", "Hesap silme"),
            ("PUT", "/api/users/{id}", "Profil güncelleme"),
            ("GET", "/api/users/admin-stats", "Admin istatistik"),
            ("POST", "/api/transfers/send-by-account", "Transfer"),
            ("POST", "/api/transfers/deposit", "Para yatırma"),
            ("GET", "/api/transactions", "İşlem listesi"),
            ("GET", "/api/fraudlogs", "Fraud logları"),
            ("GET", "/api/dailylimits", "Limitler"),
            ("POST", "/api/ai/chat", "Neon AI sohbet"),
        ],
        col_widths=[2, 6, 6.5],
    )

    add_heading(doc, "Ek B — Proje Dizin Yapısı", 1)
    add_para(
        doc,
        "Vireon/ → EntityLayer, DataAccessLayer, DtoLayer, BusinessLayer, PresentationLayer, "
        "Database/vireon_local.db, scripts/, logs/, Dockerfile, README.md, package.json"
    )

    add_heading(doc, "Ek C — Lisans", 1)
    add_para(doc, "MIT License — eğitim amaçlı final projesi (2026).")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Olusturuldu: {path}")
